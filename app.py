import os
import sqlite3
import uuid
from datetime import date, datetime, timedelta
from urllib.parse import quote
from io import BytesIO
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt
from PIL import Image

APP_DIR = Path(__file__).resolve().parent
DATA_DIR = APP_DIR / 'data'
PHOTOS_DIR = DATA_DIR / 'photos'
EXPORT_DIR = DATA_DIR / 'exports'
DB_PATH = DATA_DIR / 'buffalo_sce.db'
for p in [DATA_DIR, PHOTOS_DIR, EXPORT_DIR]:
    p.mkdir(parents=True, exist_ok=True)

st.set_page_config(page_title='Buffalo SCE Thesis App', page_icon='🦬', layout='wide')

VAGINAL_OPTIONS = ['Clear', 'Translucent', 'Cloudy']
CYCLE_OPTIONS = ['Regular', 'Irregular']
WST_OPTIONS = ['Negative', 'Slightly Positive (+)', 'Moderately Positive (++)']
PREG_OPTIONS = ['Pregnant', 'Non-pregnant', 'Pending']
VERIFICATION_OPTIONS = ['Confirmed SCE', 'Non-SCE', 'Pending']
TREATMENT_GROUPS = ['Group I - Enzyme', 'Group II - Antibiotic', 'Control / Observation']


UNIVERSITY_NAME = "SRI VENKATESWARA VETERINARY UNIVERSITY"
COLLEGE_NAME = "College of Veterinary Science, Proddatur"
DEGREE_LINE = "M.V.Sc. | Animal Reproduction Gynaecology and Obstetrics"
THESIS_TITLE = "Efficacy of Intrauterine Proteolytic Enzyme on Fertility and Assessment of Gene Expression in Buffaloes with Subclinical Endometritis"



def get_conn():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = get_conn()
    cur = conn.cursor()
    cur.execute('''
        CREATE TABLE IF NOT EXISTS cases (
            case_no TEXT PRIMARY KEY,
            village TEXT,
            record_date TEXT,
            owner_name TEXT,
            phone_no TEXT,
            buffalo_tag_no TEXT,
            breed TEXT,
            age REAL,
            parity INTEGER,
            days_postpartum INTEGER,
            bcs REAL,
            history_abortion_rp_dystocia TEXT,
            abortion_type TEXT,
            previous_ais INTEGER,
            last_heat_signs TEXT,
            heat_duration_hours REAL,
            estrous_cycle_interval TEXT,
            vaginal_discharge TEXT,
            wst_colour_desc TEXT,
            wst_result TEXT,
            blood_collection_done INTEGER,
            time_of_collection TEXT,
            epithelial_cells INTEGER,
            pmns INTEGER,
            pmn_percentage REAL,
            verification TEXT,
            treatment_group TEXT,
            treatment_date TEXT,
            treatment_notes TEXT,
            predicted_induced_heat_day9 TEXT,
            predicted_induced_heat_day10 TEXT,
            predicted_ai_48h_from_day9 TEXT,
            predicted_ai_72h_from_day9 TEXT,
            predicted_ai_48h_from_day10 TEXT,
            predicted_ai_72h_from_day10 TEXT,
            induced_heat_actual_date TEXT,
            ai_actual_date TEXT,
            induced_ai_wst TEXT,
            ai_day_blood_collection_done INTEGER,
            ai_day_time_of_collection TEXT,
            pregnancy_diagnosis TEXT,
            pregnancy_diagnosis_date TEXT,
            molecular_case_code TEXT,
            rbc_lysis_washings TEXT,
            rbc_lysis_remarks TEXT,
            bone_white INTEGER,
            lysate_storage_hrl_temp TEXT,
            lysate_storage_hrl_date TEXT,
            backup_trizol_temp TEXT,
            backup_trizol_date TEXT,
            status TEXT DEFAULT 'Open',
            created_at TEXT,
            updated_at TEXT
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS followups (
            id TEXT PRIMARY KEY,
            case_no TEXT,
            followup_date TEXT,
            followup_type TEXT,
            notes TEXT,
            induced_heat_date TEXT,
            ai_date TEXT,
            wst_result TEXT,
            blood_collection_done INTEGER,
            pregnancy_status TEXT,
            closed_this_visit INTEGER DEFAULT 0,
            created_at TEXT,
            FOREIGN KEY(case_no) REFERENCES cases(case_no)
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS photos (
            id TEXT PRIMARY KEY,
            case_no TEXT,
            visit_key TEXT,
            photo_type TEXT,
            file_path TEXT,
            caption TEXT,
            created_at TEXT,
            FOREIGN KEY(case_no) REFERENCES cases(case_no)
        )
    ''')
    conn.commit()
    conn.close()


def dt_to_str(v):
    if v is None or v == '':
        return None
    if isinstance(v, datetime):
        return v.date().isoformat()
    if isinstance(v, date):
        return v.isoformat()
    return str(v)


def save_image(uploaded_file, case_no, visit_key, photo_type):
    ext = Path(uploaded_file.name).suffix.lower() or '.jpg'
    filename = f"{case_no}_{visit_key}_{photo_type}_{uuid.uuid4().hex[:8]}{ext}"
    dest = PHOTOS_DIR / filename
    dest.write_bytes(uploaded_file.getbuffer())
    return str(dest)


def upsert_case(data):
    conn = get_conn()
    cur = conn.cursor()
    now = datetime.now().isoformat(timespec='seconds')
    exists = cur.execute('SELECT case_no FROM cases WHERE case_no=?', (data['case_no'],)).fetchone()
    cols = list(data.keys())
    if exists:
        set_clause = ', '.join([f"{c}=?" for c in cols if c != 'case_no']) + ', updated_at=?'
        values = [data[c] for c in cols if c != 'case_no'] + [now, data['case_no']]
        cur.execute(f"UPDATE cases SET {set_clause} WHERE case_no=?", values)
    else:
        insert_cols = cols + ['created_at', 'updated_at']
        placeholders = ','.join(['?'] * len(insert_cols))
        cur.execute(
            f"INSERT INTO cases ({','.join(insert_cols)}) VALUES ({placeholders})",
            [data[c] for c in cols] + [now, now]
        )
    conn.commit()
    conn.close()


def add_followup(case_no, data):
    conn = get_conn()
    cur = conn.cursor()
    fid = uuid.uuid4().hex
    cur.execute('''
        INSERT INTO followups (id, case_no, followup_date, followup_type, notes, induced_heat_date, ai_date,
                               wst_result, blood_collection_done, pregnancy_status, closed_this_visit, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        fid, case_no, data.get('followup_date'), data.get('followup_type'), data.get('notes'),
        data.get('induced_heat_date'), data.get('ai_date'), data.get('wst_result'),
        int(bool(data.get('blood_collection_done'))), data.get('pregnancy_status'),
        int(bool(data.get('closed_this_visit'))), datetime.now().isoformat(timespec='seconds')
    ))
    if data.get('closed_this_visit'):
        cur.execute("UPDATE cases SET status='Closed', updated_at=? WHERE case_no=?", (datetime.now().isoformat(timespec='seconds'), case_no))
    else:
        cur.execute("UPDATE cases SET updated_at=? WHERE case_no=?", (datetime.now().isoformat(timespec='seconds'), case_no))
    conn.commit()
    conn.close()
    return fid


def add_photo_record(case_no, visit_key, photo_type, file_path, caption=''):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute('''
        INSERT INTO photos (id, case_no, visit_key, photo_type, file_path, caption, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    ''', (uuid.uuid4().hex, case_no, visit_key, photo_type, file_path, caption, datetime.now().isoformat(timespec='seconds')))
    conn.commit()
    conn.close()


def query_df(sql, params=()):
    conn = get_conn()
    df = pd.read_sql_query(sql, conn, params=params)
    conn.close()
    return df


def fetch_case(case_no):
    conn = get_conn()
    row = conn.execute('SELECT * FROM cases WHERE case_no=?', (case_no,)).fetchone()
    conn.close()
    return dict(row) if row else None


def fetch_photos(case_no):
    return query_df('SELECT * FROM photos WHERE case_no=? ORDER BY created_at DESC', (case_no,))


def fetch_followups(case_no):
    return query_df('SELECT * FROM followups WHERE case_no=? ORDER BY followup_date DESC, created_at DESC', (case_no,))


def calc_predictions(treatment_date):
    if not treatment_date:
        return {}
    d9 = treatment_date + timedelta(days=9)
    d10 = treatment_date + timedelta(days=10)
    return {
        'predicted_induced_heat_day9': d9,
        'predicted_induced_heat_day10': d10,
        'predicted_ai_48h_from_day9': d9 + timedelta(days=2),
        'predicted_ai_72h_from_day9': d9 + timedelta(days=3),
        'predicted_ai_48h_from_day10': d10 + timedelta(days=2),
        'predicted_ai_72h_from_day10': d10 + timedelta(days=3),
    }


def pretty(v):
    if v in [None, '', 'None']:
        return '-'
    return str(v)


def google_calendar_link(title, date_iso, details=''):
    if not date_iso:
        return ''
    start = datetime.fromisoformat(str(date_iso)).strftime('%Y%m%d')
    end = (datetime.fromisoformat(str(date_iso)) + timedelta(days=1)).strftime('%Y%m%d')
    return (
        'https://calendar.google.com/calendar/render?action=TEMPLATE'
        f'&text={quote(str(title))}'
        f'&dates={start}/{end}'
        f'&details={quote(str(details))}'
    )


def build_google_calendar_links(case):
    items = [
        ('Induced heat check - Day 9', case.get('predicted_induced_heat_day9'), 'Treatment follow-up reminder'),
        ('Induced heat check - Day 10', case.get('predicted_induced_heat_day10'), 'Treatment follow-up reminder'),
        ('AI window - 48h from Day 9', case.get('predicted_ai_48h_from_day9'), 'AI timing reminder'),
        ('AI window - 72h from Day 9', case.get('predicted_ai_72h_from_day9'), 'AI timing reminder'),
        ('AI window - 48h from Day 10', case.get('predicted_ai_48h_from_day10'), 'AI timing reminder'),
        ('AI window - 72h from Day 10', case.get('predicted_ai_72h_from_day10'), 'AI timing reminder'),
    ]
    rows = []
    for label, d, desc in items:
        if d:
            title = f"{case['case_no']} - {label}"
            rows.append({'event': label, 'date': d, 'url': google_calendar_link(title, d, f"{desc} | Owner: {case.get('owner_name','-')} | Village: {case.get('village','-')}")})
    return pd.DataFrame(rows)


def add_heading_style(doc):
    styles = doc.styles
    styles['Normal'].font.name = 'Calibri'
    styles['Normal'].font.size = Pt(10.5)


def add_photo_to_doc(doc, path, caption):
    if not path or not os.path.exists(path):
        return
    try:
        img = Image.open(path)
        img.thumbnail((1200, 1200))
        temp = BytesIO()
        img.save(temp, format='JPEG')
        temp.seek(0)
        doc.add_picture(temp, width=Inches(2.2))
        if caption:
            p = doc.add_paragraph(caption)
            p.runs[0].italic = True
            p.runs[0].font.size = Pt(8.5)
    except Exception:
        doc.add_paragraph(f"[Could not render image: {path}]")


def build_case_docx(case_no):
    case = fetch_case(case_no)
    if not case:
        return None
    followups = fetch_followups(case_no)
    photos = fetch_photos(case_no)

    doc = Document()
    add_heading_style(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.6)
    sec.right_margin = Inches(0.6)

    u = doc.add_paragraph()
    u.alignment = 1
    r = u.add_run(UNIVERSITY_NAME)
    r.bold = True
    r.font.size = Pt(13)

    c = doc.add_paragraph()
    c.alignment = 1
    r = c.add_run(COLLEGE_NAME)
    r.bold = True
    r.font.size = Pt(11.5)

    d = doc.add_paragraph()
    d.alignment = 1
    d.add_run(DEGREE_LINE).italic = True

    thesis = doc.add_paragraph()
    thesis.alignment = 1
    rr = thesis.add_run('Thesis Title: ' + THESIS_TITLE)
    rr.bold = True
    rr.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = 1
    r = title.add_run('MASTER RESEARCH LOG / CASE REPORT')
    r.bold = True
    r.font.size = Pt(15)

    subtitle = doc.add_paragraph()
    subtitle.alignment = 1
    subtitle.add_run(f"Case No: {case_no}    |    Status: {pretty(case.get('status'))}").bold = True

    t = doc.add_table(rows=0, cols=2)
    t.style = 'Table Grid'
    pairs = [
        ('Date', case.get('record_date')), ('Village', case.get('village')),
        ('Owner Name', case.get('owner_name')), ('Phone No', case.get('phone_no')),
        ('Buffalo Tag No', case.get('buffalo_tag_no')), ('Breed', case.get('breed')),
        ('Age', case.get('age')), ('Parity', case.get('parity')),
        ('Days Post-Partum', case.get('days_postpartum')), ('BCS (1-5)', case.get('bcs')),
        ('History of Abortion/RP/Dystocia', case.get('history_abortion_rp_dystocia')),
        ('Abortion Type', case.get('abortion_type')), ('Previous AIs', case.get('previous_ais')),
        ('Nature of Last Heat Signs', case.get('last_heat_signs')),
        ('Heat Duration (hrs)', case.get('heat_duration_hours')),
        ('Estrous Cycle Interval', case.get('estrous_cycle_interval')),
        ('Vaginal Discharge', case.get('vaginal_discharge')),
        ('WST Colour Description', case.get('wst_colour_desc')),
        ('WST Result', case.get('wst_result')), ('Blood Collection Done', 'Yes' if case.get('blood_collection_done') else 'No'),
        ('Time of Collection', case.get('time_of_collection')), ('Total Epithelial Cells', case.get('epithelial_cells')),
        ('Total PMNs', case.get('pmns')), ('PMN Percentage', case.get('pmn_percentage')),
        ('Verification', case.get('verification')), ('Treatment Group', case.get('treatment_group')),
        ('Treatment Date', case.get('treatment_date')), ('Treatment Notes', case.get('treatment_notes')),
        ('Predicted Induced Heat - Day 9', case.get('predicted_induced_heat_day9')),
        ('Predicted Induced Heat - Day 10', case.get('predicted_induced_heat_day10')),
        ('Predicted AI 48h from Day 9', case.get('predicted_ai_48h_from_day9')),
        ('Predicted AI 72h from Day 9', case.get('predicted_ai_72h_from_day9')),
        ('Predicted AI 48h from Day 10', case.get('predicted_ai_48h_from_day10')),
        ('Predicted AI 72h from Day 10', case.get('predicted_ai_72h_from_day10')),
        ('Actual Induced Heat Date', case.get('induced_heat_actual_date')),
        ('Actual AI Date', case.get('ai_actual_date')), ('Induced AI Day WST', case.get('induced_ai_wst')),
        ('AI Day Blood Collection', 'Yes' if case.get('ai_day_blood_collection_done') else 'No'),
        ('AI Day Time of Collection', case.get('ai_day_time_of_collection')),
        ('Pregnancy Diagnosis', case.get('pregnancy_diagnosis')),
        ('Pregnancy Diagnosis Date', case.get('pregnancy_diagnosis_date')),
        ('Molecular Case Code', case.get('molecular_case_code')),
        ('RBC Lysis Washings', case.get('rbc_lysis_washings')),
        ('RBC Lysis Remarks', case.get('rbc_lysis_remarks')),
        ('Bone White', 'Yes' if case.get('bone_white') else 'No'),
        ('HRL + β-ME Temp', case.get('lysate_storage_hrl_temp')),
        ('HRL + β-ME Date', case.get('lysate_storage_hrl_date')),
        ('TRIzol Temp', case.get('backup_trizol_temp')),
        ('TRIzol Date', case.get('backup_trizol_date')),
    ]
    for a, b in pairs:
        row = t.add_row().cells
        row[0].text = str(a)
        row[1].text = pretty(b)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Follow-up Log').bold = True
    if followups.empty:
        doc.add_paragraph('No follow-up entries yet.')
    else:
        ft = doc.add_table(rows=1, cols=7)
        ft.style = 'Table Grid'
        hdr = ft.rows[0].cells
        hdr[0].text = 'Date'
        hdr[1].text = 'Type'
        hdr[2].text = 'Induced Heat'
        hdr[3].text = 'AI Date'
        hdr[4].text = 'WST'
        hdr[5].text = 'Pregnancy'
        hdr[6].text = 'Notes'
        for _, r in followups.iterrows():
            rr = ft.add_row().cells
            rr[0].text = pretty(r['followup_date'])
            rr[1].text = pretty(r['followup_type'])
            rr[2].text = pretty(r['induced_heat_date'])
            rr[3].text = pretty(r['ai_date'])
            rr[4].text = pretty(r['wst_result'])
            rr[5].text = pretty(r['pregnancy_status'])
            rr[6].text = pretty(r['notes'])

    doc.add_page_break()
    h = doc.add_paragraph()
    h.add_run('Photo Documentation').bold = True
    if photos.empty:
        doc.add_paragraph('No photos uploaded yet.')
    else:
        grouped = photos.groupby(['visit_key', 'photo_type'], dropna=False)
        for (visit_key, photo_type), g in grouped:
            p = doc.add_paragraph()
            p.add_run(f"{visit_key} - {photo_type}").bold = True
            for _, row in g.head(6).iterrows():
                add_photo_to_doc(doc, row['file_path'], row['caption'])

    out_path = EXPORT_DIR / f"{case_no}_case_report.docx"
    doc.save(out_path)
    return out_path


def build_master_csv():
    df = query_df('SELECT * FROM cases ORDER BY updated_at DESC')
    if df.empty:
        return None
    out = EXPORT_DIR / 'buffalo_sce_master_cases.csv'
    df.to_csv(out, index=False)
    return out


def build_due_reminders_csv():
    today = date.today().isoformat()
    df = query_df('''
        SELECT case_no, owner_name, village, phone_no, status,
               predicted_induced_heat_day9, predicted_induced_heat_day10,
               predicted_ai_48h_from_day9, predicted_ai_72h_from_day9,
               predicted_ai_48h_from_day10, predicted_ai_72h_from_day10,
               induced_heat_actual_date, ai_actual_date
        FROM cases
        WHERE status='Open'
        ORDER BY predicted_induced_heat_day9 ASC
    ''')
    if df.empty:
        return None
    out = EXPORT_DIR / f'due_reminders_{today}.csv'
    df.to_csv(out, index=False)
    return out


def build_ics_for_case(case):
    def event_block(start_date, title, desc):
        if not start_date:
            return ''
        dt = datetime.fromisoformat(start_date).strftime('%Y%m%d')
        uid = uuid.uuid4().hex
        return f"BEGIN:VEVENT\nUID:{uid}\nDTSTAMP:{datetime.utcnow().strftime('%Y%m%dT%H%M%SZ')}\nDTSTART;VALUE=DATE:{dt}\nSUMMARY:{title}\nDESCRIPTION:{desc}\nEND:VEVENT\n"

    body = 'BEGIN:VCALENDAR\nVERSION:2.0\nPRODID:-//Buffalo SCE Thesis App//EN\n'
    body += event_block(case.get('predicted_induced_heat_day9'), f"{case['case_no']} induced heat check (Day 9)", 'Follow-up from treatment date')
    body += event_block(case.get('predicted_induced_heat_day10'), f"{case['case_no']} induced heat check (Day 10)", 'Follow-up from treatment date')
    body += event_block(case.get('predicted_ai_48h_from_day9'), f"{case['case_no']} AI window (48h from Day 9)", 'AI reminder')
    body += event_block(case.get('predicted_ai_72h_from_day9'), f"{case['case_no']} AI window (72h from Day 9)", 'AI reminder')
    body += event_block(case.get('predicted_ai_48h_from_day10'), f"{case['case_no']} AI window (48h from Day 10)", 'AI reminder')
    body += event_block(case.get('predicted_ai_72h_from_day10'), f"{case['case_no']} AI window (72h from Day 10)", 'AI reminder')
    body += 'END:VCALENDAR\n'
    out = EXPORT_DIR / f"{case['case_no']}_reminders.ics"
    out.write_text(body, encoding='utf-8')
    return out


def show_due_cases():
    df = query_df('''
        SELECT case_no, owner_name, village, phone_no, status,
               predicted_induced_heat_day9, predicted_induced_heat_day10,
               predicted_ai_48h_from_day9, predicted_ai_72h_from_day9,
               predicted_ai_48h_from_day10, predicted_ai_72h_from_day10
        FROM cases WHERE status='Open'
        ORDER BY predicted_induced_heat_day9 ASC
    ''')
    if df.empty:
        st.info('No open cases yet.')
        return
    st.dataframe(df, use_container_width=True)


def safe_int(v):
    try:
        return int(v)
    except Exception:
        return None


def safe_float(v):
    try:
        return float(v)
    except Exception:
        return None


def normalize_page_label(label):
    mapping = {
        'New / Edit Case': 'New / Edit Case',
        'Reopen Case / Follow-up': 'Reopen Case / Follow-up',
        'Dashboard & Reminders': 'Dashboard & Reminders',
        'Export Reports': 'Export Reports',
    }
    return mapping.get(label, 'Dashboard & Reminders')


def build_case_action_link(case_no, page='Reopen Case / Follow-up', reminder_date=None, reminder_type=None):
    base = '?page=' + quote(page) + '&case_no=' + quote(str(case_no))
    if reminder_date:
        base += '&reminder_date=' + quote(str(reminder_date))
    if reminder_type:
        base += '&reminder_type=' + quote(str(reminder_type))
    return base


def reminder_rows_for_case(case):
    rows = []
    if not case or case.get('status') == 'Closed':
        return rows
    mapping = [
        ('Induced heat check - Day 9', case.get('predicted_induced_heat_day9')),
        ('Induced heat check - Day 10', case.get('predicted_induced_heat_day10')),
        ('AI window - 48h from Day 9', case.get('predicted_ai_48h_from_day9')),
        ('AI window - 72h from Day 9', case.get('predicted_ai_72h_from_day9')),
        ('AI window - 48h from Day 10', case.get('predicted_ai_48h_from_day10')),
        ('AI window - 72h from Day 10', case.get('predicted_ai_72h_from_day10')),
    ]
    done_induced = bool(case.get('induced_heat_actual_date'))
    done_ai = bool(case.get('ai_actual_date'))
    for reminder_type, d in mapping:
        if not d:
            continue
        done = False
        if 'Induced heat' in reminder_type and done_induced:
            done = True
        if 'AI window' in reminder_type and done_ai:
            done = True
        rows.append({
            'case_no': case.get('case_no'),
            'owner_name': case.get('owner_name'),
            'village': case.get('village'),
            'phone_no': case.get('phone_no'),
            'status': case.get('status'),
            'reminder_type': reminder_type,
            'reminder_date': d,
            'done': done,
            'open_link': build_case_action_link(case.get('case_no'), reminder_date=d, reminder_type=reminder_type)
        })
    return rows


def build_internal_reminders_df(include_completed=False):
    cases = query_df('SELECT * FROM cases ORDER BY updated_at DESC')
    rows = []
    if cases.empty:
        return pd.DataFrame()
    for _, r in cases.iterrows():
        rows.extend(reminder_rows_for_case(r.to_dict()))
    df = pd.DataFrame(rows)
    if df.empty:
        return df
    df['reminder_date'] = pd.to_datetime(df['reminder_date'])
    today = pd.Timestamp(date.today())
    def classify(v):
        if v.date() < today.date():
            return 'Overdue'
        if v.date() == today.date():
            return 'Today'
        return 'Upcoming'
    df['bucket'] = df['reminder_date'].apply(classify)
    df['days_from_today'] = (df['reminder_date'].dt.date - today.date()).apply(lambda x: x.days)
    if not include_completed:
        df = df[df['done'] == False]
    return df.sort_values(['reminder_date', 'case_no']).reset_index(drop=True)


def render_reminder_cards(df, limit=None):
    if df.empty:
        st.info('No pending reminders.')
        return
    view = df.head(limit) if limit else df
    for _, row in view.iterrows():
        label = f"{row['bucket']} | {row['reminder_date'].date()} | {row['reminder_type']}"
        st.markdown(f"**{row['case_no']}** — {pretty(row['owner_name'])} ({pretty(row['village'])})  \n{label}  \n[Open this case now]({row['open_link']})")
        st.divider()


init_db()

st.title('🦬 Buffalo SCE Thesis App')
st.caption('Mobile-friendly case entry, reopen same case for follow-up, internal reminder dashboard with direct case-action links, Google Calendar reminder links, .ics export, and thesis-style Word report generation.')

query_page = normalize_page_label(st.query_params.get('page', 'Dashboard & Reminders'))
pages = ['New / Edit Case', 'Reopen Case / Follow-up', 'Dashboard & Reminders', 'Export Reports']
default_idx = pages.index(query_page) if query_page in pages else 2
page = st.sidebar.radio('Open section', pages, index=default_idx)
if st.query_params.get('page') != page:
    st.query_params['page'] = page
selected_case_from_link = st.query_params.get('case_no', '')
linked_reminder_date = st.query_params.get('reminder_date', '')
linked_reminder_type = st.query_params.get('reminder_type', '')

if page == 'New / Edit Case':
    st.subheader('Create or update a case')
    with st.form('case_form', clear_on_submit=False):
        c1, c2, c3, c4 = st.columns(4)
        case_no = c1.text_input('Case No *')
        village = c2.text_input('Village')
        record_date = c3.date_input('Record Date', value=date.today())
        owner_name = c4.text_input('Owner Name')

        c1, c2, c3, c4 = st.columns(4)
        phone_no = c1.text_input('Phone No')
        buffalo_tag_no = c2.text_input('Buffalo Tag No')
        breed = c3.text_input('Breed')
        age = c4.number_input('Age', min_value=0.0, step=0.5)

        c1, c2, c3, c4 = st.columns(4)
        parity = c1.number_input('Parity', min_value=0, step=1)
        days_postpartum = c2.number_input('Days Post-Partum', min_value=0, step=1)
        bcs = c3.number_input('BCS (1-5)', min_value=0.0, max_value=5.0, step=0.1)
        previous_ais = c4.number_input('No. of Previous AIs', min_value=0, step=1)

        history_abortion = st.radio('History of Abortion / RP / Dystocia', ['No', 'Yes'], horizontal=True)
        abortion_type = st.text_input('Type (only if yes)')
        c1, c2, c3 = st.columns(3)
        last_heat_signs = c1.text_input('Nature of Last Heat Signs')
        heat_duration_hours = c2.number_input('Duration of Heat (hrs)', min_value=0.0, step=0.5)
        estrous_cycle_interval = c3.selectbox('Estrous Cycle Interval', CYCLE_OPTIONS)
        vaginal_discharge = st.selectbox('Vaginal Discharge', VAGINAL_OPTIONS)

        st.markdown('### Field action and cytology')
        c1, c2, c3 = st.columns(3)
        wst_colour_desc = c1.text_area('White Side Test Colour Description')
        wst_result = c2.selectbox('WST Result', WST_OPTIONS)
        blood_collection_done = c3.checkbox('Blood collection in EDTA done')
        time_of_collection = st.text_input('Time of Collection')

        c1, c2, c3 = st.columns(3)
        epithelial_cells = c1.number_input('Total Epithelial Cells', min_value=0, step=1)
        pmns = c2.number_input('Total PMNs', min_value=0, step=1)
        pmn_percentage = c3.number_input('PMN Percentage', min_value=0.0, max_value=100.0, step=0.1)
        verification = st.selectbox('Verification', VERIFICATION_OPTIONS)

        st.markdown('### Treatment and prediction')
        c1, c2 = st.columns(2)
        treatment_group = c1.selectbox('Treatment Group', TREATMENT_GROUPS)
        treatment_date = c2.date_input('Treatment Date', value=date.today())
        treatment_notes = st.text_area('Treatment Notes')

        preds = calc_predictions(treatment_date)
        st.info(
            f"Predicted induced heat: {preds['predicted_induced_heat_day9']} or {preds['predicted_induced_heat_day10']}\n\n"
            f"Predicted AI windows:\n"
            f"• From day 9: {preds['predicted_ai_48h_from_day9']} to {preds['predicted_ai_72h_from_day9']}\n"
            f"• From day 10: {preds['predicted_ai_48h_from_day10']} to {preds['predicted_ai_72h_from_day10']}"
        )

        st.markdown('### Clinical follow-up now available')
        c1, c2, c3 = st.columns(3)
        induced_heat_actual_date = c1.date_input('Actual Induced Heat Date', value=None)
        ai_actual_date = c2.date_input('Actual AI Date', value=None)
        induced_ai_wst = c3.selectbox('Induced AI Day WST', ['-', 'Neg', 'Pos'])
        c1, c2, c3 = st.columns(3)
        ai_day_blood_collection_done = c1.checkbox('AI day blood collection done')
        ai_day_time_of_collection = c2.text_input('AI day time of collection')
        pregnancy_diagnosis = c3.selectbox('Pregnancy Diagnosis', PREG_OPTIONS)
        pregnancy_diagnosis_date = st.date_input('Pregnancy Diagnosis Date', value=None)

        st.markdown('### Molecular metadata')
        c1, c2 = st.columns(2)
        molecular_case_code = c1.text_input('Code Given / Molecular Case Code')
        rbc_lysis_washings = c2.text_input('RBC Lysis - No. of Washings')
        c1, c2 = st.columns(2)
        rbc_lysis_remarks = c1.text_input('RBC Lysis Remarks')
        bone_white = c2.checkbox('Bone white')
        c1, c2, c3, c4 = st.columns(4)
        lysate_storage_hrl_temp = c1.text_input('HRL + β-ME Temp')
        lysate_storage_hrl_date = c2.date_input('HRL + β-ME Date', value=None)
        backup_trizol_temp = c3.text_input('TRIzol Temp')
        backup_trizol_date = c4.date_input('TRIzol Date', value=None)

        st.markdown('### Photos - initial visit')
        uterine_discharge_photo = st.file_uploader('Uterine discharge photo', type=['jpg', 'jpeg', 'png'], accept_multiple_files=False)
        wst_before = st.file_uploader('White Side Test photo - before', type=['jpg', 'jpeg', 'png'], accept_multiple_files=False)
        wst_after = st.file_uploader('White Side Test photo - after', type=['jpg', 'jpeg', 'png'], accept_multiple_files=False)
        cytology_photo = st.file_uploader('Cytology photo', type=['jpg', 'jpeg', 'png'], accept_multiple_files=False)
        extra_initial = st.file_uploader('Additional initial visit photos', type=['jpg', 'jpeg', 'png'], accept_multiple_files=True)

        close_case_now = st.checkbox('Close case now')
        submitted = st.form_submit_button('Save case')

    if submitted:
        if not case_no.strip():
            st.error('Case No is required.')
        else:
            payload = {
                'case_no': case_no.strip(),
                'village': village, 'record_date': dt_to_str(record_date), 'owner_name': owner_name,
                'phone_no': phone_no, 'buffalo_tag_no': buffalo_tag_no, 'breed': breed,
                'age': safe_float(age), 'parity': safe_int(parity), 'days_postpartum': safe_int(days_postpartum),
                'bcs': safe_float(bcs), 'history_abortion_rp_dystocia': history_abortion,
                'abortion_type': abortion_type if history_abortion == 'Yes' else '',
                'previous_ais': safe_int(previous_ais), 'last_heat_signs': last_heat_signs,
                'heat_duration_hours': safe_float(heat_duration_hours), 'estrous_cycle_interval': estrous_cycle_interval,
                'vaginal_discharge': vaginal_discharge, 'wst_colour_desc': wst_colour_desc,
                'wst_result': wst_result, 'blood_collection_done': int(bool(blood_collection_done)),
                'time_of_collection': time_of_collection, 'epithelial_cells': safe_int(epithelial_cells),
                'pmns': safe_int(pmns), 'pmn_percentage': safe_float(pmn_percentage), 'verification': verification,
                'treatment_group': treatment_group, 'treatment_date': dt_to_str(treatment_date),
                'treatment_notes': treatment_notes,
                'predicted_induced_heat_day9': dt_to_str(preds['predicted_induced_heat_day9']),
                'predicted_induced_heat_day10': dt_to_str(preds['predicted_induced_heat_day10']),
                'predicted_ai_48h_from_day9': dt_to_str(preds['predicted_ai_48h_from_day9']),
                'predicted_ai_72h_from_day9': dt_to_str(preds['predicted_ai_72h_from_day9']),
                'predicted_ai_48h_from_day10': dt_to_str(preds['predicted_ai_48h_from_day10']),
                'predicted_ai_72h_from_day10': dt_to_str(preds['predicted_ai_72h_from_day10']),
                'induced_heat_actual_date': dt_to_str(induced_heat_actual_date), 'ai_actual_date': dt_to_str(ai_actual_date),
                'induced_ai_wst': induced_ai_wst if induced_ai_wst != '-' else '',
                'ai_day_blood_collection_done': int(bool(ai_day_blood_collection_done)),
                'ai_day_time_of_collection': ai_day_time_of_collection,
                'pregnancy_diagnosis': pregnancy_diagnosis, 'pregnancy_diagnosis_date': dt_to_str(pregnancy_diagnosis_date),
                'molecular_case_code': molecular_case_code, 'rbc_lysis_washings': rbc_lysis_washings,
                'rbc_lysis_remarks': rbc_lysis_remarks, 'bone_white': int(bool(bone_white)),
                'lysate_storage_hrl_temp': lysate_storage_hrl_temp, 'lysate_storage_hrl_date': dt_to_str(lysate_storage_hrl_date),
                'backup_trizol_temp': backup_trizol_temp, 'backup_trizol_date': dt_to_str(backup_trizol_date),
                'status': 'Closed' if close_case_now else 'Open'
            }
            upsert_case(payload)
            visit_key = 'initial'
            upload_map = {
                'uterine_discharge': uterine_discharge_photo,
                'wst_before': wst_before,
                'wst_after': wst_after,
                'cytology': cytology_photo,
            }
            for label, f in upload_map.items():
                if f is not None:
                    path = save_image(f, case_no.strip(), visit_key, label)
                    add_photo_record(case_no.strip(), visit_key, label, path, label.replace('_', ' ').title())
            if extra_initial:
                for i, f in enumerate(extra_initial, start=1):
                    path = save_image(f, case_no.strip(), visit_key, f'extra_{i}')
                    add_photo_record(case_no.strip(), visit_key, f'extra_{i}', path, f'Extra initial photo {i}')
            st.success(f'Case {case_no} saved successfully.')
            st.markdown(f"[Open this case for follow-up now]({build_case_action_link(case_no.strip())})")
            case = fetch_case(case_no.strip())
            ics_path = build_ics_for_case(case)
            with open(ics_path, 'rb') as fh:
                st.download_button('Download reminders (.ics)', data=fh.read(), file_name=ics_path.name, mime='text/calendar')
            links_df = build_google_calendar_links(case)
            if not links_df.empty:
                st.markdown('### Add reminders to Google Calendar')
                for _, row in links_df.iterrows():
                    st.markdown(f"- [{row['event']} ({row['date']})]({row['url']})")

elif page == 'Reopen Case / Follow-up':
    st.subheader('Reopen existing case and add follow-up')
    case_list = query_df('SELECT case_no, owner_name, village, status, updated_at FROM cases ORDER BY updated_at DESC')
    if case_list.empty:
        st.info('No cases saved yet.')
    else:
        case_options = case_list['case_no'].tolist()
        pre_idx = case_options.index(selected_case_from_link) if selected_case_from_link in case_options else 0
        selected = st.selectbox('Select Case No', case_options, index=pre_idx)
        case = fetch_case(selected)
        if case:
            if linked_reminder_type or linked_reminder_date:
                st.info(f"Opened from reminder: {linked_reminder_type or 'Case reminder'} | Date: {linked_reminder_date or '-'}")
            c1, c2, c3, c4 = st.columns(4)
            c1.metric('Owner', pretty(case.get('owner_name')))
            c2.metric('Village', pretty(case.get('village')))
            c3.metric('Status', pretty(case.get('status')))
            c4.metric('PMN %', pretty(case.get('pmn_percentage')))
            st.write('Predicted reminder dates')
            st.write({
                'Induced heat (day 9)': case.get('predicted_induced_heat_day9'),
                'Induced heat (day 10)': case.get('predicted_induced_heat_day10'),
                'AI 48h from day 9': case.get('predicted_ai_48h_from_day9'),
                'AI 72h from day 9': case.get('predicted_ai_72h_from_day9'),
                'AI 48h from day 10': case.get('predicted_ai_48h_from_day10'),
                'AI 72h from day 10': case.get('predicted_ai_72h_from_day10'),
            })
            links_df = build_google_calendar_links(case)
            if not links_df.empty:
                st.markdown('### Google Calendar links')
                for _, row in links_df.iterrows():
                    st.markdown(f"- [{row['event']} ({row['date']})]({row['url']})")
            with st.form('followup_form'):
                c1, c2 = st.columns(2)
                followup_date = c1.date_input('Follow-up Date', value=date.today())
                followup_type = c2.selectbox('Follow-up Type', ['Induced Heat Check', 'AI Visit', 'Pregnancy Diagnosis', 'Lab Update', 'General Review'])
                c1, c2, c3 = st.columns(3)
                induced_heat_date = c1.date_input('Induced Heat Date', value=None)
                ai_date = c2.date_input('AI Date', value=None)
                followup_wst = c3.selectbox('WST Result', ['-', 'Neg', 'Pos', 'Slightly Positive (+)', 'Moderately Positive (++)'])
                c1, c2 = st.columns(2)
                blood_done = c1.checkbox('Blood collection done')
                pregnancy_status = c2.selectbox('Pregnancy Status', PREG_OPTIONS)
                notes = st.text_area('Follow-up Notes')
                st.markdown('### Photos - follow-up visit')
                fu_uterine_discharge = st.file_uploader('Uterine discharge photo (follow-up)', type=['jpg', 'jpeg', 'png'], accept_multiple_files=False)
                fu_wst_before = st.file_uploader('WST photo - before (follow-up)', type=['jpg', 'jpeg', 'png'], accept_multiple_files=False)
                fu_wst_after = st.file_uploader('WST photo - after (follow-up)', type=['jpg', 'jpeg', 'png'], accept_multiple_files=False)
                fu_cytology = st.file_uploader('Cytology photo (follow-up)', type=['jpg', 'jpeg', 'png'], accept_multiple_files=False)
                fu_extra = st.file_uploader('Additional follow-up photos', type=['jpg', 'jpeg', 'png'], accept_multiple_files=True)
                close_case = st.checkbox('Close case after this follow-up')
                save_fu = st.form_submit_button('Save follow-up')
            if save_fu:
                data = {
                    'followup_date': dt_to_str(followup_date),
                    'followup_type': followup_type,
                    'notes': notes,
                    'induced_heat_date': dt_to_str(induced_heat_date),
                    'ai_date': dt_to_str(ai_date),
                    'wst_result': '' if followup_wst == '-' else followup_wst,
                    'blood_collection_done': blood_done,
                    'pregnancy_status': pregnancy_status,
                    'closed_this_visit': close_case,
                }
                visit_key = f"followup_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                add_followup(selected, data)
                upload_map = {
                    'uterine_discharge': fu_uterine_discharge,
                    'wst_before': fu_wst_before,
                    'wst_after': fu_wst_after,
                    'cytology': fu_cytology,
                }
                for label, f in upload_map.items():
                    if f is not None:
                        path = save_image(f, selected, visit_key, label)
                        add_photo_record(selected, visit_key, label, path, f'{label.replace("_", " ").title()} - follow-up')
                if fu_extra:
                    for i, f in enumerate(fu_extra, start=1):
                        path = save_image(f, selected, visit_key, f'extra_{i}')
                        add_photo_record(selected, visit_key, f'extra_{i}', path, f'Extra follow-up photo {i}')
                st.success('Follow-up saved.')

            st.markdown('### Existing follow-up log')
            st.dataframe(fetch_followups(selected), use_container_width=True)
            st.markdown('### Uploaded photos')
            photos = fetch_photos(selected)
            if photos.empty:
                st.info('No photos uploaded for this case yet.')
            else:
                st.dataframe(photos[['visit_key', 'photo_type', 'caption', 'created_at']], use_container_width=True)

elif page == 'Dashboard & Reminders':
    st.subheader('Dashboard')
    total_cases = query_df('SELECT COUNT(*) AS n FROM cases')['n'].iloc[0]
    open_cases = query_df("SELECT COUNT(*) AS n FROM cases WHERE status='Open'")['n'].iloc[0]
    closed_cases = query_df("SELECT COUNT(*) AS n FROM cases WHERE status='Closed'")['n'].iloc[0]
    preg_cases = query_df("SELECT COUNT(*) AS n FROM cases WHERE pregnancy_diagnosis='Pregnant'")['n'].iloc[0]
    c1, c2, c3, c4 = st.columns(4)
    c1.metric('Total Cases', int(total_cases))
    c2.metric('Open Cases', int(open_cases))
    c3.metric('Closed Cases', int(closed_cases))
    c4.metric('Pregnant', int(preg_cases))
    st.markdown('### Internal reminder to-do list')
    reminder_df = build_internal_reminders_df(include_completed=False)
    if reminder_df.empty:
        st.info('No pending reminders.')
    else:
        c1, c2, c3 = st.columns(3)
        c1.metric('Due today', int((reminder_df['bucket'] == 'Today').sum()))
        c2.metric('Overdue', int((reminder_df['bucket'] == 'Overdue').sum()))
        c3.metric('Upcoming', int((reminder_df['bucket'] == 'Upcoming').sum()))
        render_reminder_cards(reminder_df, limit=20)
        st.dataframe(reminder_df[['case_no', 'owner_name', 'village', 'reminder_type', 'reminder_date', 'bucket', 'open_link']], use_container_width=True)
    st.markdown('### Open case master list')
    show_due_cases()
    master_csv = build_master_csv()
    if master_csv and master_csv.exists():
        with open(master_csv, 'rb') as fh:
            st.download_button('Download master CSV', data=fh.read(), file_name=master_csv.name, mime='text/csv')
    due_csv = build_due_reminders_csv()
    if due_csv and due_csv.exists():
        with open(due_csv, 'rb') as fh:
            st.download_button('Download due reminders CSV', data=fh.read(), file_name=due_csv.name, mime='text/csv')
    st.info('Each reminder now has a direct action link that opens the same case in the follow-up screen. You can still use Google Calendar links or the .ics file if you also want calendar alerts.')

elif page == 'Export Reports':
    st.subheader('Export case-wise Word report')
    case_list = query_df('SELECT case_no, owner_name, status, updated_at FROM cases ORDER BY updated_at DESC')
    if case_list.empty:
        st.info('No cases available.')
    else:
        selected = st.selectbox('Select case for Word export', case_list['case_no'].tolist(), key='export_case')
        if st.button('Generate Word report'):
            out = build_case_docx(selected)
            if out and out.exists():
                with open(out, 'rb') as fh:
                    st.download_button('Download Word report', data=fh.read(), file_name=out.name, mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                st.success('Word report generated.')
        case = fetch_case(selected)
        if case:
            ics_path = build_ics_for_case(case)
            with open(ics_path, 'rb') as fh:
                st.download_button('Download this case reminder file (.ics)', data=fh.read(), file_name=ics_path.name, mime='text/calendar')
            links_df = build_google_calendar_links(case)
            if not links_df.empty:
                st.markdown('### Google Calendar links for this case')
                for _, row in links_df.iterrows():
                    st.markdown(f"- [{row['event']} ({row['date']})]({row['url']})")
